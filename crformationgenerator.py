import streamlit as st
import pandas as pd
from docx import Document
import os
import tempfile
from zipfile import ZipFile
import random
import re
from collections import defaultdict

st.set_page_config(page_title="Générateur de QCM par session", layout="centered")
st.title("📄 Générateur de QCM par session (figé ou aléatoire)")

# Fonction de remplacement des balises
def remplacer_placeholders(paragraph, replacements):
    """
    Parcourt chaque run d'un paragraphe et remplace les clés du dictionnaire `replacements`
    par leurs valeurs respectives, sans écraser le formatage local (gras, italique, etc.).
    """
    for key, val in replacements.items():
        if key in paragraph.text:
            for run in paragraph.runs:
                if key in run.text:
                    run.text = run.text.replace(key, val)

# Fonction pour itérer sur tous les paragraphes (y compris ceux dans les tableaux)
def iter_all_paragraphs(doc):
    """
    Génère tous les objets Paragraph d'un Document python‐docx, 
    qu'ils soient à la racine ou à l'intérieur de cellules de tableaux.
    """
    # Paragraphes à la racine
    for para in doc.paragraphs:
        yield para

    # Paragraphes à l'intérieur des cellules de chaque tableau
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para

# Définition des réponses "positives" pour chaque groupe (prioritaires en mode aléatoire)
POSITIVE_OPTIONS = {
    "satisfaction": ["Très satisfait", "Satisfait"],
    "motivation": ["Très motivés", "Motivés"],
    "assiduite": ["Très motivés", "Motivés"],
    "homogeneite": ["Oui"],
    "questions": ["Toutes les questions", "A peu près toutes"],
    "adaptation": ["Non"],   # Même si on va finalement forcer "Non", on garde ici pour la logique par défaut
    "suivi": ["Non concerné"]
}

# Définition des options possibles pour chaque groupe de cases à cocher
CHECKBOX_GROUPS = {
    "satisfaction": ["Très satisfait", "Satisfait", "Moyennement satisfait", "Insatisfait", "Non satisfait"],
    "motivation": ["Très motivés", "Motivés", "Pas motivés"],
    "assiduite": ["Très motivés", "Motivés", "Pas motivés"],
    "homogeneite": ["Oui", "Non"],
    "questions": [
        "Toutes les questions",
        "A peu près toutes",
        "Il y a quelques sujets sur lesquels je n'avais pas les réponses",
        "Je n'ai pas pu répondre à la majorité des questions"
    ],
    "adaptation": ["Oui", "Non"],
    "suivi": ["Oui", "Non", "Non concerné"]
}

# Étape 1 : Importer les fichiers (Excel + Word)
with st.expander("Etape 1 : Importer les fichiers", expanded=True):
    excel_file = st.file_uploader("Fichier Excel des participants", type="xlsx")
    word_file = st.file_uploader("Modèle Word du compte rendu", type="docx")

# Si les deux fichiers sont fournis, on entre dans la génération
if excel_file and word_file:
    # Lecture du fichier Excel en DataFrame
    df = pd.read_excel(excel_file)
    # On enlève d'éventuels espaces superflus dans les noms de colonnes
    df.columns = df.columns.str.strip()

    # Vérification que toutes les colonnes obligatoires sont présentes
    required_columns = ["session", "formateur", "formation", "nb d'heure", "Nom", "Prénom"]
    if not set(required_columns).issubset(df.columns):
        st.error(f"Colonnes manquantes dans le fichier Excel. Colonnes requises : {required_columns}")
        st.info(f"Colonnes disponibles : {list(df.columns)}")
    else:
        # On groupe les participants par "session"
        sessions = df.groupby("session")
        # Dictionnaire des réponses figées choisies par l'utilisateur
        reponses_figees = {}

        # Étape 2 : Si l’utilisateur souhaite figer certaines réponses, on affiche un checkbox + selectbox
        st.markdown("### Etape 2 : Choisir les réponses à figer (facultatif)")
        for groupe, options in CHECKBOX_GROUPS.items():
            figer = st.checkbox(f"Figer la réponse pour : {groupe}", key=f"figer_{groupe}")
            if figer:
                # Si l’utilisateur coche “figer” pour ce groupe, on lui propose toutes les options à figer
                choix = st.selectbox(f"Choix figé pour {groupe}", options, key=f"choix_{groupe}")
                reponses_figees[groupe] = choix

        # Champs libres pour “Avis & pistes d'amélioration” et “Autres observations”
        pistes = st.text_area("Avis & pistes d'amélioration :", key="pistes")
        observations = st.text_area("Autres observations :", key="obs")

        # Bouton pour lancer la génération des comptes rendus
        if st.button("🚀 Générer les comptes rendus"):
            # On crée un dossier temporaire pour stocker les .docx générés et l'archive ZIP
            with tempfile.TemporaryDirectory() as tmpdir:
                zip_path = os.path.join(tmpdir, "QCM_Sessions.zip")
                with ZipFile(zip_path, 'w') as zipf:
                    # Pour chaque session (clé `session_id`, DataFrame `participants`)
                    for session_id, participants in sessions:
                        # On ouvre une nouvelle instance du modèle Word pour cette session
                        doc = Document(word_file)
                        first = participants.iloc[0]  # On récupère la première ligne pour certaines infos

                        # Préparation des remplacements pour tous les placeholders standards
                        replacements = {
                            "{{nom}}": str(first["Nom"]),
                            "{{prénom}}": str(first["Prénom"]),
                            "{{formateur}}": f"{first['Prénom']} {first['Nom']}",
                            "{{ref_session}}": str(session_id),
                            "{{formation_dispensee}}": str(first["formation"]),
                            "{{duree_formation}}": str(first["nb d'heure"]),
                            "{{nb_participants}}": str(len(participants))
                        }

                        # 1) On remplace tous les placeholders de type {{nom}}, {{prénom}}, etc.
                        for para in iter_all_paragraphs(doc):
                            remplacer_placeholders(para, replacements)

                        # 2) On repère tous les paragraphes qui contiennent le placeholder "{{checkbox}}"
                        checkbox_paras = []
                        for para in iter_all_paragraphs(doc):
                            if "{{checkbox}}" in para.text:
                                # On normalise un peu le texte pour pouvoir matcher les options
                                texte = re.sub(r'\s+', ' ', para.text).strip()
                                for groupe, options in CHECKBOX_GROUPS.items():
                                    for opt in options:
                                        # Si le texte du paragraphe contient exactement l'option (délimitée par \b)
                                        if re.search(rf"\b{re.escape(opt)}\b", texte):
                                            checkbox_paras.append((groupe, opt, para))
                                            break
                                    else:
                                        # Aucun opt ne matché dans ce groupe, on passe au groupe suivant
                                        continue
                                    # On a trouvé un match dans ce groupe, on sort de la boucle d'options
                                    break

                        # 3) On regroupe les paragraphes par groupe (ex : “satisfaction”, “adaptation”, etc.)
                        group_to_paras = defaultdict(list)
                        for groupe, opt, para in checkbox_paras:
                            group_to_paras[groupe].append((opt, para))

                        # 4) Pour chaque groupe, on décide quelle case doit être cochée définitivement
                        for groupe, paras in group_to_paras.items():
                            # Liste des options réellement présentes dans le modèle pour ce groupe
                            options_presentes = [opt for opt, _ in paras]

                            # —————————————————————————————————————————————
                            # LOGIQUE DE FIGEAGE DÉFINITIF POUR "adaptation" ET "suivi"
                            # —————————————————————————————————————————————
                            if groupe == "adaptation" and "Non" in options_presentes:
                                # On force toujours “Non” pour le groupe “adaptation”
                                option_choisie = "Non"
                            elif groupe == "suivi" and "Non concerné" in options_presentes:
                                # On force toujours “Non concerné” pour le groupe “suivi”
                                option_choisie = "Non concerné"

                            # —————————————————————————————————————————————
                            # SINON, SI L’UTILISATEUR A DEMANDÉ UN FIGEAGE VIA L’INTERFACE
                            # —————————————————————————————————————————————
                            elif groupe in reponses_figees:
                                option_choisie = reponses_figees[groupe]

                            # —————————————————————————————————————————————
                            # SINON, LOGIQUE “ALÉATOIRE” (PRIORITÉ AUX RÉPONSES POSITIVES)
                            # —————————————————————————————————————————————
                            else:
                                # On prend d’abord toutes les options marquées “positives”
                                positives_disponibles = [
                                    opt
                                    for opt in options_presentes
                                    if groupe in POSITIVE_OPTIONS and opt in POSITIVE_OPTIONS[groupe]
                                ]
                                if positives_disponibles:
                                    option_choisie = random.choice(positives_disponibles)
                                else:
                                    # Si aucune “positive” présente, on tire n’importe quelle option
                                    option_choisie = (
                                        random.choice(options_presentes)
                                        if options_presentes
                                        else None
                                    )

                            # 5) On parcourt chaque paragraphe du groupe et on remplace {{checkbox}}
                            #     par “☑” si c’est l’option à cocher, sinon “☐”.
                            if option_choisie:
                                for opt, para in paras:
                                    for run in para.runs:
                                        if "{{checkbox}}" in run.text:
                                            run.text = run.text.replace(
                                                "{{checkbox}}",
                                                "☑" if opt == option_choisie else "☐"
                                            )

                        # 6) On ajoute ensuite, à la fin du document, les sections libres
                        doc.add_paragraph("\nAvis & pistes d'amélioration :\n" + pistes)
                        doc.add_paragraph("\nAutres observations :\n" + observations)

                        # 7) Enregistrement du document .docx pour cette session
                        filename = f"Compte_Rendu_{session_id}.docx"
                        path = os.path.join(tmpdir, filename)
                        doc.save(path)
                        zipf.write(path, arcname=filename)

                # 8) Une fois toutes les sessions traitées, on propose le téléchargement du ZIP
                with open(zip_path, "rb") as f:
                    st.success("Comptes rendus générés avec succès !")
                    st.download_button(
                        "📅 Télécharger l'archive ZIP",
                        data=f,
                        file_name="QCM_Sessions.zip",
                        mime="application/zip"
                    )
